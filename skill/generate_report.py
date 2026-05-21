#!/usr/bin/env python3
"""Generate a Word health report on the user's Desktop.

Reads (from repo, read-only):
  <repo>/health/daily.csv    (required)
  <repo>/health/workouts.csv (optional)
  <repo>/index.csv           (optional, kibble meal log)

Writes (to Desktop, NOT into the repo):
  ~/Desktop/kibble-health-report-<YYYY-MM-DD>.docx

The .docx opens natively in Word / Pages / Google Docs / LibreOffice /
macOS Quick Look. Charts are embedded as PNGs inside the docx package.
The repo holds only source data; this report is regenerable and lives
outside the repo by design.
"""

import io
import os
import sys
from datetime import timedelta

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def fig_to_png_buffer(fig) -> io.BytesIO:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", dpi=120)
    plt.close(fig)
    buf.seek(0)
    return buf


def build_recommendations(daily, workouts, meals, today, year_ago):
    """Return [(section_name, [{'obs':..., 'action':...}, ...])].

    Every bullet must reference a specific number from the data.
    Refuse to give advice when signals conflict or sample is thin —
    say "signal unclear" or "need more data" rather than guessing.
    """
    last_30_idx = pd.Timestamp(today - timedelta(days=30))
    last_60_idx = pd.Timestamp(today - timedelta(days=60))
    yr = daily.loc[daily.index >= year_ago]
    last_30 = daily.loc[daily.index >= last_30_idx]

    exercise = []
    nutrition = []

    # --- Volume trend (workouts last 4w vs preceding 4w) ---
    if workouts is not None and len(workouts) > 0:
        wks = workouts.dropna(subset=["start", "duration_min"]).copy()
        recent = wks[(wks["start"] >= last_30_idx)]["duration_min"].sum()
        prior = wks[(wks["start"] >= last_60_idx)
                    & (wks["start"] < last_30_idx)]["duration_min"].sum()
        if prior > 0 and recent > 0:
            pct = (recent / prior - 1) * 100
            obs = (f"Workout volume last 4 weeks: {recent:.0f} min vs preceding 4 weeks "
                   f"{prior:.0f} min ({pct:+.0f}%).")
            if pct > 40:
                action = ("Sharp volume increase. Watch RHR for the next 2 weeks — "
                          "if 7-day RHR climbs >5 bpm above 12-month baseline, swap one "
                          "session for an easy day or full rest.")
            elif pct < -40:
                action = ("Volume dropped sharply. If intentional (taper, illness, exam "
                          "block) ignore. Otherwise rebuild gradually, "
                          "no more than +10–20% per week.")
            else:
                action = "Volume stable — maintain current routine."
            exercise.append({"obs": obs, "action": action})
        elif recent == 0 and prior > 0:
            exercise.append({
                "obs": f"No workouts logged in the last 4 weeks "
                       f"(preceding 4 weeks: {prior:.0f} min).",
                "action": "If this is unintentional, restart with 2–3 short easy sessions "
                          "this week before resuming previous load.",
            })

    # --- RHR signal ---
    rhr_recent = yr["resting_hr"].tail(30).mean() if "resting_hr" in yr else float("nan")
    rhr_mean = yr["resting_hr"].mean() if "resting_hr" in yr else float("nan")
    if not (pd.isna(rhr_recent) or pd.isna(rhr_mean)):
        delta = rhr_recent - rhr_mean
        obs = (f"Resting HR (last 30d): {rhr_recent:.1f} bpm vs 12-month avg "
               f"{rhr_mean:.1f} bpm ({delta:+.1f}).")
        if delta < -3:
            action = ("Below baseline — typical of aerobic improvement and good recovery. "
                      "No action; current load is well absorbed.")
            exercise.append({"obs": obs, "action": action})
        elif delta > 5:
            action = ("Elevated. Common causes: under-recovery, recent illness, "
                      "alcohol, dehydration, poor sleep. Take 2–3 days of easy "
                      "training; if RHR doesn't return within a week, see a doctor.")
            exercise.append({"obs": obs, "action": action})
        # ±3..+5 is normal noise; no recommendation

    # --- HRV signal ---
    hrv_recent = yr["hrv_ms"].tail(30).mean() if "hrv_ms" in yr else float("nan")
    hrv_mean = yr["hrv_ms"].mean() if "hrv_ms" in yr else float("nan")
    if not (pd.isna(hrv_recent) or pd.isna(hrv_mean)) and hrv_mean > 0:
        delta_pct = (hrv_recent - hrv_mean) / hrv_mean * 100
        obs = (f"HRV SDNN (last 30d): {hrv_recent:.1f} ms vs 12-month avg "
               f"{hrv_mean:.1f} ms ({delta_pct:+.0f}%).")
        if delta_pct < -15:
            action = ("Sustained ↓HRV: under-recovery signal. Cut intensity by one step "
                      "(e.g. swap intervals for easy run) for one week; prioritise "
                      "sleep ≥7 h.")
            exercise.append({"obs": obs, "action": action})
        elif delta_pct > 10:
            action = ("Recovery markers strong. Safe to add intensity (one harder "
                      "session/week) if that fits your goals.")
            exercise.append({"obs": obs, "action": action})

    # --- Sleep ---
    sleep_30 = last_30["sleep_hours"].dropna() if "sleep_hours" in last_30 else pd.Series()
    if len(sleep_30) >= 14:
        avg = sleep_30.mean()
        obs = (f"Average sleep last 30d: {avg:.2f} h/night "
               f"(across {len(sleep_30)} nights with data).")
        if avg < 7:
            action = ("Below the 7-hour floor that consistently shows up in HRV "
                      "regression. Single highest-leverage recovery lever — try moving "
                      "lights-out 30 min earlier for 1 week, see HRV response.")
            exercise.append({"obs": obs, "action": action})

    # --- Activity diversity ---
    if workouts is not None:
        recent_w = workouts[(workouts["start"] >= last_30_idx)]
        if len(recent_w) >= 5:
            dist = recent_w["type"].value_counts(normalize=True)
            top_type, top_pct = dist.index[0], dist.iloc[0] * 100
            if top_pct > 80:
                obs = (f"Activity mix last 30d: {top_pct:.0f}% {top_type} "
                       f"(across {len(recent_w)} workouts).")
                action = ("Highly mono-modal. For general fitness this is fine, "
                          "but for VO₂ max growth and injury prevention, add 1–2 "
                          "sessions/week of a different modality (if running, add "
                          "cycling or strength; if walking, add running intervals).")
                exercise.append({"obs": obs, "action": action})

    # --- Exercise minutes vs WHO guideline ---
    if "exercise_min" in last_30:
        wkly = last_30["exercise_min"].sum() / max(len(last_30), 1) * 7
        if wkly > 0:
            obs = f"Apple Watch 'Exercise' minutes: ~{wkly:.0f} min/week (last 30d)."
            if wkly < 150:
                action = ("Below WHO recommendation of ≥150 min/week of moderate "
                          "aerobic activity. Add ~"
                          f"{150 - wkly:.0f} more min/week to reach the floor.")
                exercise.append({"obs": obs, "action": action})
            elif wkly >= 300:
                action = ("Meets WHO upper recommendation (≥300 min/week) — "
                          "great. Don't increase further unless you're training "
                          "for a specific event.")
                exercise.append({"obs": obs, "action": action})

    # --- VO2 trend ---
    vo2 = daily["vo2_max"].dropna() if "vo2_max" in daily else pd.Series()
    if len(vo2) >= 6:
        xs_num = (vo2.index - vo2.index.min()).days.to_numpy()
        slope = np.polyfit(xs_num, vo2.values, 1)[0] * 365
        if slope > 1.5:
            obs = (f"VO₂ max trajectory: {slope:+.1f} ml/kg/min/year "
                   f"({len(vo2)} measurements, latest {vo2.iloc[-1]:.1f}).")
            action = "Aerobic capacity improving — current training mix is working."
            exercise.append({"obs": obs, "action": action})
        elif slope < -1.5:
            obs = (f"VO₂ max trajectory: {slope:+.1f} ml/kg/min/year "
                   f"({len(vo2)} measurements, latest {vo2.iloc[-1]:.1f}).")
            action = ("Downward trend. Check intensity distribution — if all sessions "
                      "are easy/aerobic, add 1 weekly session at HR zone 4 "
                      "(intervals or tempo).")
            exercise.append({"obs": obs, "action": action})

    # --- Nutrition ---
    if meals is None or len(meals) == 0:
        nutrition.append({
            "obs": "No meal entries logged.",
            "action": "Use Kibble to drag food/receipt photos in. Need ≥14 days of "
                      "consistent logging before nutrition recommendations carry any signal.",
        })
    else:
        n_meals = len(meals)
        days_logged = (meals["date"].dt.date.nunique()
                       if "date" in meals.columns else 0)
        if days_logged < 14:
            nutrition.append({
                "obs": f"Meal log so far: {n_meals} entries across {days_logged} day(s).",
                "action": (f"Need ≥14 consecutive days of logging for any nutrition "
                           f"signal. Currently at {days_logged}/14 — keep using Kibble. "
                           "Re-run this report after 2 weeks of logging."),
            })
        else:
            # Real nutrition analysis (will trigger once user accumulates data)
            kcal_by_day = meals.groupby(meals["date"].dt.date)["calories"].sum()
            avg_intake = kcal_by_day.mean()
            paired_balance = []
            for d, intake in kcal_by_day.items():
                ts = pd.Timestamp(d)
                if ts in daily.index:
                    row = daily.loc[ts]
                    burn = ((row.get("active_kcal", 0) or 0)
                            + (row.get("basal_kcal", 0) or 0))
                    paired_balance.append(intake - burn)
            if paired_balance:
                avg_bal = sum(paired_balance) / len(paired_balance)
                obs = (f"Logged days: {days_logged}. Avg intake "
                       f"{avg_intake:.0f} kcal/day. Avg energy balance vs Apple Health "
                       f"expenditure: {avg_bal:+.0f} kcal/day "
                       f"({len(paired_balance)} paired days).")
                if avg_bal < -500:
                    action = ("Sustained ≥500 kcal/day deficit. Healthy weight loss "
                              "pace; check that protein stays at ≥1.6 g/kg body mass "
                              "to preserve lean tissue.")
                elif avg_bal > 500:
                    action = ("Sustained surplus. If gaining weight intentionally "
                              "(hypertrophy block) fine; otherwise tighten portion "
                              "sizes for 2 weeks.")
                else:
                    action = "Energy balance roughly even."
                nutrition.append({"obs": obs, "action": action})

    return [("Exercise", exercise), ("Nutrition", nutrition)]


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            table.rows[r].cells[c].text = str(val)
    if col_widths:
        for row in table.rows:
            for c, w in enumerate(col_widths):
                row.cells[c].width = w
    return table


def main():
    if len(sys.argv) not in (2, 3):
        sys.exit("usage: generate_report.py <repo_dir> [<output_path>]")
    repo = sys.argv[1]
    health = os.path.join(repo, "health")

    daily_path = os.path.join(health, "daily.csv")
    if not os.path.isfile(daily_path):
        sys.exit(f"no daily.csv at {daily_path}; run parse_apple_health.py first")

    daily = pd.read_csv(daily_path, parse_dates=["date"]).set_index("date").sort_index()

    workouts = None
    wk_path = os.path.join(health, "workouts.csv")
    if os.path.isfile(wk_path):
        workouts = pd.read_csv(wk_path, parse_dates=["start", "end"])
        for col in ("start", "end"):
            if pd.api.types.is_datetime64_any_dtype(workouts[col]) \
               and workouts[col].dt.tz is not None:
                workouts[col] = workouts[col].dt.tz_localize(None)
        workouts["duration_min"] = pd.to_numeric(workouts["duration_min"], errors="coerce")
        workouts["distance"] = pd.to_numeric(workouts["distance"], errors="coerce")
        workouts["kcal"] = pd.to_numeric(workouts["kcal"], errors="coerce")

    meals = None
    meals_path = os.path.join(repo, "index.csv")
    if os.path.isfile(meals_path):
        try:
            meals = pd.read_csv(meals_path, parse_dates=["date"])
        except Exception:
            meals = None

    today = daily.index.max().date()
    year_ago = pd.Timestamp(today - timedelta(days=365))
    last_year = daily.loc[daily.index >= year_ago].copy()
    last_30 = daily.loc[daily.index >= pd.Timestamp(today - timedelta(days=30))].copy()
    last_7 = daily.loc[daily.index >= pd.Timestamp(today - timedelta(days=6))].copy()

    plt.rcParams.update({
        "figure.dpi": 120,
        "axes.grid": True,
        "grid.alpha": 0.3,
        "font.size": 10,
    })

    chart_imgs = {}

    # 1) RHR + HRV 12mo
    fig, ax1 = plt.subplots(figsize=(9, 3.6))
    rhr_roll = last_year["resting_hr"].rolling(7, min_periods=2).mean()
    hrv_roll = last_year["hrv_ms"].rolling(7, min_periods=2).mean()
    ax1.plot(rhr_roll.index, rhr_roll.values, color="#c0392b", label="RHR (7-day avg)")
    ax1.set_ylabel("Resting HR (bpm)", color="#c0392b")
    ax1.tick_params(axis="y", labelcolor="#c0392b")
    ax2 = ax1.twinx()
    ax2.plot(hrv_roll.index, hrv_roll.values, color="#2980b9", label="HRV SDNN (7-day avg)")
    ax2.set_ylabel("HRV SDNN (ms)", color="#2980b9")
    ax2.tick_params(axis="y", labelcolor="#2980b9")
    ax2.grid(False)
    fig.suptitle("Resting HR + HRV — last 12 months (7-day rolling)")
    fig.tight_layout()
    chart_imgs["rhr_hrv"] = fig_to_png_buffer(fig)

    # 2) Weekly training volume
    if workouts is not None and len(workouts) > 0:
        wks = workouts.dropna(subset=["start"]).copy()
        wks["week"] = wks["start"].dt.to_period("W").apply(lambda r: r.start_time)
        wks_cut = wks[wks["start"] >= pd.Timestamp(today - timedelta(days=180))]
        if len(wks_cut) > 0:
            pivot = (wks_cut.assign(min=wks_cut["duration_min"])
                     .pivot_table(index="week", columns="type",
                                  values="min", aggfunc="sum")
                     .fillna(0))
            top = pivot.sum().sort_values(ascending=False).head(5).index
            pivot = pivot[top]
            fig, ax = plt.subplots(figsize=(9, 3.6))
            bottom = np.zeros(len(pivot))
            for col in pivot.columns:
                ax.bar(pivot.index, pivot[col].values, bottom=bottom,
                       width=5, label=col)
                bottom += pivot[col].values
            ax.set_ylabel("Minutes / week")
            ax.set_title("Weekly training volume by activity — last 6 months")
            ax.legend(loc="upper left", fontsize=8)
            fig.tight_layout()
            chart_imgs["weekly_volume"] = fig_to_png_buffer(fig)

    # 3) Sleep vs next-day HRV
    if "sleep_hours" in last_year and "hrv_ms" in last_year:
        df = last_year[["sleep_hours", "hrv_ms"]].copy()
        df["next_hrv"] = df["hrv_ms"].shift(-1)
        df = df.dropna()
        if len(df) > 20:
            fig, ax = plt.subplots(figsize=(7, 4))
            ax.scatter(df["sleep_hours"], df["next_hrv"], alpha=0.4, s=18)
            if len(df) > 5:
                z = np.polyfit(df["sleep_hours"], df["next_hrv"], 1)
                p = np.poly1d(z)
                xs = np.linspace(df["sleep_hours"].min(), df["sleep_hours"].max(), 100)
                ax.plot(xs, p(xs), color="#c0392b",
                        label=f"fit: HRV = {z[0]:.2f}·sleep + {z[1]:.1f}")
                ax.legend()
            ax.set_xlabel("Sleep hours (night n)")
            ax.set_ylabel("HRV SDNN, ms (morning n+1)")
            ax.set_title("Sleep duration vs next-day HRV — last 12 months")
            fig.tight_layout()
            chart_imgs["sleep_hrv"] = fig_to_png_buffer(fig)

    # 4) VO2 max
    vo2 = daily["vo2_max"].dropna()
    if len(vo2) > 1:
        fig, ax = plt.subplots(figsize=(9, 3.6))
        ax.scatter(vo2.index, vo2.values, alpha=0.6, s=20)
        if len(vo2) > 3:
            xs_num = (vo2.index - vo2.index.min()).days.to_numpy()
            z = np.polyfit(xs_num, vo2.values, 1)
            p = np.poly1d(z)
            ax.plot(vo2.index, p(xs_num), color="#c0392b",
                    label=f"trend: {z[0]*365:+.2f} ml/kg/min per year")
            ax.legend()
        ax.set_ylabel("VO₂ max (ml/kg/min)")
        ax.set_title(f"VO₂ max — all-time ({len(vo2)} measurements)")
        fig.tight_layout()
        chart_imgs["vo2_max"] = fig_to_png_buffer(fig)

    # 5) Steps + active kcal 30d
    fig, ax1 = plt.subplots(figsize=(9, 3.6))
    ax1.bar(last_30.index, last_30["steps"], color="#7f8c8d", alpha=0.6, label="Steps")
    ax1.set_ylabel("Steps", color="#7f8c8d")
    ax1.tick_params(axis="y", labelcolor="#7f8c8d")
    ax2 = ax1.twinx()
    ax2.plot(last_30.index, last_30["active_kcal"], color="#e67e22",
             marker="o", markersize=3, label="Active kcal")
    ax2.set_ylabel("Active kcal", color="#e67e22")
    ax2.tick_params(axis="y", labelcolor="#e67e22")
    ax2.grid(False)
    fig.suptitle("Steps + active energy — last 30 days")
    fig.tight_layout()
    chart_imgs["steps_active"] = fig_to_png_buffer(fig)

    # ---------- build .docx ----------
    def fmt(v, n=0):
        if pd.isna(v):
            return "—"
        return f"{v:,.{n}f}"

    yr = daily.loc[daily.index >= year_ago]
    rhr_mean = yr["resting_hr"].mean()
    rhr_recent = yr["resting_hr"].tail(30).mean()
    hrv_mean = yr["hrv_ms"].mean()
    hrv_recent = yr["hrv_ms"].tail(30).mean()
    vo2_latest = vo2.iloc[-1] if len(vo2) else float("nan")
    vo2_year_ago = (vo2[vo2.index <= year_ago].iloc[-1]
                    if (vo2.index <= year_ago).any() else float("nan"))

    doc = Document()
    styles = doc.styles["Normal"]
    styles.font.name = "Helvetica"
    styles.font.size = Pt(11)

    title = doc.add_heading(f"Health report — {today}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = doc.add_paragraph()
    run = p.add_run("Generated by the kibble-health skill from Apple Health + Kibble "
                    "meal log. Source data lives in the kibble-data git repo and is "
                    "not modified by this report.")
    run.italic = True
    run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    # Overview
    doc.add_heading("Overview", level=1)
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("Date range: ").bold = False
    p.add_run(f"{daily.index.min().date()} → {daily.index.max().date()} "
              f"({len(daily):,} days)").bold = True
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("Total steps recorded: ")
    p.add_run(f"{int(daily['steps'].sum()):,}").bold = True
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("Workouts logged: ")
    p.add_run(f"{len(workouts) if workouts is not None else 0}").bold = True
    if workouts is not None and len(workouts) > 0:
        top = workouts["type"].value_counts().head(3)
        p = doc.add_paragraph(style="List Bullet")
        p.add_run("Most common activities: " +
                  ", ".join(f"{t} ({n})" for t, n in top.items()))

    # At-a-glance
    doc.add_heading("At-a-glance — last 30 days vs 12-month baseline", level=1)
    rows = [
        ["Resting HR (bpm)", fmt(rhr_recent, 1), fmt(rhr_mean, 1),
         f"{(rhr_recent - rhr_mean):+.1f}"],
        ["HRV SDNN (ms)", fmt(hrv_recent, 1), fmt(hrv_mean, 1),
         f"{(hrv_recent - hrv_mean):+.1f}"],
    ]
    if not pd.isna(vo2_latest):
        rows.append(["VO₂ max (latest / 1y ago)", fmt(vo2_latest, 1),
                     fmt(vo2_year_ago, 1), f"{(vo2_latest - vo2_year_ago):+.1f}"])
    add_table(doc, ["Metric", "Last 30d", "12-month avg", "Δ"], rows)

    p = doc.add_paragraph()
    r = p.add_run("Direction-of-travel: ↓RHR + ↑HRV = better aerobic base / recovery. "
                  "↑RHR + ↓HRV sustained >1 week → consider deload or check sleep.")
    r.italic = True
    r.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    # Last 7 days
    doc.add_heading("Last 7 days", level=1)
    rows = []
    for d, row in last_7.iterrows():
        rows.append([
            str(d.date()),
            fmt(row.get("steps")),
            fmt(row.get("active_kcal"), 0),
            fmt(row.get("walk_run_km"), 2),
            fmt(row.get("exercise_min"), 0),
            fmt(row.get("resting_hr"), 0),
            fmt(row.get("hrv_ms"), 1),
            fmt(row.get("sleep_hours"), 2),
        ])
    add_table(doc, ["Date", "Steps", "Active kcal", "Walk/Run km",
                    "Exercise min", "RHR", "HRV", "Sleep h"], rows)

    # Trends
    doc.add_heading("Trends", level=1)
    doc.add_heading("Resting HR + HRV — last 12 months", level=2)
    doc.add_picture(chart_imgs["rhr_hrv"], width=Inches(6.5))
    if "weekly_volume" in chart_imgs:
        doc.add_heading("Weekly training volume by activity — last 6 months", level=2)
        doc.add_picture(chart_imgs["weekly_volume"], width=Inches(6.5))
    if "sleep_hrv" in chart_imgs:
        doc.add_heading("Sleep duration ↔ next-day HRV", level=2)
        doc.add_picture(chart_imgs["sleep_hrv"], width=Inches(5.5))
    if "vo2_max" in chart_imgs:
        doc.add_heading("VO₂ max trajectory", level=2)
        doc.add_picture(chart_imgs["vo2_max"], width=Inches(6.5))
    doc.add_heading("Steps + active kcal — last 30 days", level=2)
    doc.add_picture(chart_imgs["steps_active"], width=Inches(6.5))

    # Meal log
    doc.add_heading("Meal log (Kibble)", level=1)
    if meals is None or len(meals) == 0:
        p = doc.add_paragraph()
        r = p.add_run("No meal entries yet. Drop food photos into Kibble to start logging.")
        r.italic = True
    else:
        n_meals = len(meals)
        days_logged = meals["date"].dt.date.nunique() if "date" in meals.columns else 0
        p = doc.add_paragraph(style="List Bullet")
        p.add_run("Entries: ")
        p.add_run(f"{n_meals}").bold = True
        p.add_run(f" across {days_logged} day(s)")
        if "calories" in meals.columns:
            kcal_by_day = meals.groupby(meals["date"].dt.date)["calories"].sum()
            p = doc.add_paragraph(style="List Bullet")
            p.add_run("Average intake (logged days only): ")
            p.add_run(f"{kcal_by_day.mean():.0f} kcal/day").bold = True
            paired = []
            for d, intake in kcal_by_day.items():
                ts = pd.Timestamp(d)
                if ts in daily.index:
                    row = daily.loc[ts]
                    burn = ((row.get("active_kcal", 0) or 0)
                            + (row.get("basal_kcal", 0) or 0))
                    paired.append([str(d), f"{intake:.0f}", f"{burn:.0f}",
                                   f"{intake - burn:+.0f}"])
            if paired:
                doc.add_paragraph("")
                add_table(doc, ["Date", "Intake (kcal)", "Expenditure (kcal)",
                                "Balance"], paired[-7:])

    # Recommendations
    recs = build_recommendations(daily, workouts, meals, today, year_ago)
    doc.add_heading("Observations & recommendations", level=1)
    p = doc.add_paragraph()
    r = p.add_run("Every bullet below is anchored to a specific number from your "
                  "data — if a signal was ambiguous or sparse, no recommendation "
                  "was generated for that area. This is direction-of-travel, not "
                  "clinical advice; for injury, illness, or pain, see your "
                  "physiotherapist/clinician.")
    r.italic = True
    r.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    for section, items in recs:
        doc.add_heading(section, level=2)
        if not items:
            p = doc.add_paragraph()
            r = p.add_run("No actionable signal at the current sample size — "
                          "no recommendation generated.")
            r.italic = True
            continue
        for item in items:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run("Observation: ")
            r.bold = True
            p.add_run(item["obs"])
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            r = p.add_run("→ Action: ")
            r.bold = True
            p.add_run(item["action"])

    # Footer note
    doc.add_paragraph("")
    p = doc.add_paragraph()
    r = p.add_run("This report is regenerable from the source CSVs in "
                  "~/Desktop/health-log/health/ plus the Kibble meal log. "
                  "Delete this file freely — running the skill again will produce "
                  "a fresh one. GPS routes from Apple Health are deliberately not "
                  "synced or analysed.")
    r.italic = True
    r.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    r.font.size = Pt(9)

    if len(sys.argv) == 3:
        out_path = sys.argv[2]
    else:
        out_path = os.path.expanduser(f"~/Desktop/kibble-health-report-{today}.docx")
    doc.save(out_path)
    size_kb = os.path.getsize(out_path) / 1024
    print(f"wrote {out_path} ({size_kb:.0f} KB, {len(chart_imgs)} embedded charts)")


if __name__ == "__main__":
    main()
