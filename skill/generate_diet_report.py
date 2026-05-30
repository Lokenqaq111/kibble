#!/usr/bin/env python3
"""Diet report generator for the kibble health-log.

Usage:
    generate_diet_report.py <repo_path>

Reads <repo>/index.csv, aggregates per day, and writes a Word report to the
Desktop (NEVER into the repo -- derived artefact). The report contains:
  * overview stats (days logged, breakfast coverage, fast-food / sweet counts)
  * a per-day macro table with auto notes
  * three trend line-charts (calories, fat, protein) with red threshold lines;
    days crossing the line are marked red
  * a data-driven "signals" section -- every bullet references a real number,
    no platitudes (same rule as generate_report.py's build_recommendations()).

Thresholds and keyword sets are configurable at the top. The qualitative
written interpretation is intentionally NOT hardcoded here; Claude layers that
on top after running the script (see SKILL.md, Diet report mode).
"""
import csv, os, sys, datetime, tempfile
from collections import defaultdict
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

# --- tunables --------------------------------------------------------------
# kind: "ceiling" -> a day OVER the line is flagged red (calories, fat).
#       "floor"   -> a day UNDER the line is flagged red (protein minimum).
THRESHOLDS = {
    "cal": {"label": "Calories (kcal/day)", "line": 2000, "kind": "ceiling"},
    "f":   {"label": "Fat (g/day)",         "line": 70,   "kind": "ceiling"},
    "p":   {"label": "Protein (g/day)",     "line": 70,   "kind": "floor"},
}
FAT_NOTE_G   = 100   # per-day fat above this gets a table note
HIGH_CAL_DAY = 2200  # "true intake" hint ceiling used only in the caption

# label keyword buckets (food_label is lowercase English in index.csv)
FAST = ("mcdonalds", "kfc", "mcnuggets", "burger king", "dicos", "tastien")
SWEET = ("ice cream", "wafer", "cake", "chips", "pastry", "cola", "chocolate",
         "cookies", "wife cake", "cheese pastry", "swiss roll", "smoothie",
         "milk drink", "latte", "vienna")
FRIED = ("fried", "cutlet", "nuggets", "mcnuggets", "pastry rolls", "tonkatsu")
VEG = ("salad", "vegetable", "veg", "greens", "broccoli", "spinach", "kale",
       "bok choy", "choy", "lettuce")
# ---------------------------------------------------------------------------


def num(x):
    x = (x or "").strip()
    try:
        return float(x)
    except ValueError:
        return 0.0


def load(repo):
    csv_path = os.path.join(repo, "index.csv")
    if not os.path.isfile(csv_path):
        sys.exit(f"no index.csv at {csv_path} -- run Process mode first")
    rows = []
    with open(csv_path, newline="", encoding="utf-8") as f:
        for r in csv.DictReader(f):
            rows.append(r)
    if not rows:
        sys.exit("index.csv is empty -- nothing to report")
    return rows


def aggregate(rows):
    days = defaultdict(lambda: {"cal": 0.0, "p": 0.0, "c": 0.0, "f": 0.0,
                                "items": 0, "meals": set(),
                                "fast": 0, "sweet": 0, "fried": 0, "veg": 0})
    for r in rows:
        lbl = r["food_label"].lower()
        e = days[r["date"]]
        e["cal"] += num(r["calories"]); e["p"] += num(r["protein_g"])
        e["c"] += num(r["carb_g"]); e["f"] += num(r["fat_g"])
        e["items"] += 1; e["meals"].add(r["meal_type"])
        if any(k in lbl for k in FAST):  e["fast"] += 1
        if any(k in lbl for k in SWEET): e["sweet"] += 1
        if any(k in lbl for k in FRIED): e["fried"] += 1
        if any(k in lbl for k in VEG):   e["veg"] += 1
    return days


def make_charts(days, complete, outdir):
    xs = [d[5:] for d in complete]            # MM-DD labels
    out = []
    for key, cfg in THRESHOLDS.items():
        ys = [days[d][key] for d in complete]
        thr, kind = cfg["line"], cfg["kind"]
        flag = [(y > thr) if kind == "ceiling" else (y < thr) for y in ys]
        fig, ax = plt.subplots(figsize=(7.2, 2.7), dpi=130)
        ax.plot(xs, ys, color="#3a6ea5", lw=1.8, marker="o", ms=5,
                mfc="#3a6ea5", mec="#3a6ea5", zorder=2)
        rx = [xs[i] for i in range(len(xs)) if flag[i]]
        ry = [ys[i] for i in range(len(xs)) if flag[i]]
        ax.scatter(rx, ry, color="#d62828", s=70, zorder=3)
        for x, y in zip(rx, ry):
            ax.annotate(f"{y:.0f}", (x, y), textcoords="offset points",
                        xytext=(0, 8), ha="center", fontsize=8,
                        color="#d62828", fontweight="bold")
        ax.axhline(thr, color="#d62828", ls="--", lw=1.3)
        word = "ceiling" if kind == "ceiling" else "floor (min)"
        ax.text(0.995, thr, f" {word}: {thr}", color="#d62828", fontsize=8,
                va="bottom" if kind == "ceiling" else "top", ha="right",
                transform=ax.get_yaxis_transform())
        ax.set_title(cfg["label"], fontsize=10, fontweight="bold", loc="left")
        ax.grid(axis="y", ls=":", alpha=0.45)
        ax.margins(x=0.03); ax.set_ylim(bottom=0)
        for s in ("top", "right"):
            ax.spines[s].set_visible(False)
        fig.tight_layout()
        p = os.path.join(outdir, f"chart_{key}.png")
        fig.savefig(p); plt.close(fig)
        out.append((key, cfg, flag, p))
    return out


def build_signals(days, dates, complete, rows):
    """Deterministic observations. Each item references a real number; if a
    signal is in the noise band, NO bullet is generated for it (no platitudes).
    Returns (good, issues) lists of (title, body) tuples."""
    n = len(dates); nc = len(complete) or 1
    avg = lambda k: sum(days[d][k] for d in complete) / nc
    good, issues = [], []

    bf_days = sum(1 for d in dates if "breakfast" in days[d]["meals"])
    fast = sum(days[d]["fast"] for d in dates)
    sweet = sum(days[d]["sweet"] for d in dates)
    veg_days = sum(1 for d in dates if days[d]["veg"] > 0)
    fat_over = [d[5:] for d in complete if days[d]["f"] > THRESHOLDS["f"]["line"]]
    pro_under = [d[5:] for d in complete if days[d]["p"] < THRESHOLDS["p"]["line"]]
    cal_over = [d[5:] for d in complete if days[d]["cal"] > THRESHOLDS["cal"]["line"]]

    # strengths
    if avg("p") >= THRESHOLDS["p"]["line"]:
        good.append(("蛋白质达标", f"完整日蛋白均值约 {avg('p'):.0f} g，达到/超过下限 "
                                   f"{THRESHOLDS['p']['line']} g。"))
    if len(rows) >= n * 2:
        good.append(("记录密度高", f"{n} 天共 {len(rows)} 条记录，平均每天 "
                                   f"{len(rows)/n:.1f} 条，记录习惯稳定。"))

    # issues (only when the number crosses a band)
    if bf_days <= n / 2:
        issues.append(("早餐缺失", f"{n} 天里仅 {bf_days} 天记录早餐，易形成"
                                   "「空腹到中午再暴一顿」的模式。"))
    if fast >= max(2, n / 3):
        issues.append(("快餐频繁", f"快餐类记录 {fast} 次，约每 {n/fast:.1f} 天一次。"))
    if sweet >= n:
        issues.append(("甜食/油炸零食偏多", f"含糖或油炸零食 {sweet} 次，"
                                            f"平均每天 {sweet/n:.1f} 次。"))
    if veg_days <= n / 3:
        issues.append(("蔬菜/纤维不足", f"仅 {veg_days}/{n} 天的记录里出现蔬菜成分。"))
    if fat_over:
        issues.append(("脂肪超上限", f"{len(fat_over)} 天超过 {THRESHOLDS['f']['line']} g："
                                     + "、".join(fat_over) + "。"))
    if cal_over:
        issues.append(("热量超上限", f"{len(cal_over)} 天超过 "
                                     f"{THRESHOLDS['cal']['line']} kcal："
                                     + "、".join(cal_over) + "。"))
    if pro_under:
        issues.append(("蛋白不达标日", f"{len(pro_under)} 天低于下限 "
                                       f"{THRESHOLDS['p']['line']} g："
                                       + "、".join(pro_under) + "。"))
    return good, issues


def build_report(repo):
    rows = load(repo)
    days = aggregate(rows)
    dates = sorted(days)
    complete = [d for d in dates
                if len(days[d]["meals"]) >= 2 or days[d]["items"] >= 2]
    nc = len(complete) or 1
    avg = lambda k: sum(days[d][k] for d in complete) / nc
    bf_days = sum(1 for d in dates if "breakfast" in days[d]["meals"])
    fast = sum(days[d]["fast"] for d in dates)
    sweet = sum(days[d]["sweet"] for d in dates)
    today = datetime.date.today().isoformat()
    out = os.path.expanduser(f"~/Desktop/kibble-diet-report-{today}.docx")

    doc = Document()
    st = doc.styles["Normal"]; st.font.name = "Helvetica"; st.font.size = Pt(10.5)

    doc.add_heading("饮食记录评估报告", 0)
    sub = doc.add_paragraph(f"数据区间：{dates[0]} → {dates[-1]}　|　"
                            f"生成日期：{today}　|　来源：kibble health-log")
    sub.runs[0].italic = True; sub.runs[0].font.size = Pt(9)

    # 1. overview
    doc.add_heading("一、记录概况", 1)
    p = doc.add_paragraph()
    p.add_run(f"记录跨度 {len(dates)} 天，共 {len(rows)} 条食物记录。").bold = True
    doc.add_paragraph(f"其中 {len(complete)} 天记录较完整；"
                      f"{len(dates)-len(complete)} 天为残缺（仅单餐）。"
                      f"全程 {bf_days} 天记录了早餐。", style="List Bullet")
    doc.add_paragraph(f"快餐类记录约 {fast} 次；含糖/油炸零食类约 {sweet} 次。",
                      style="List Bullet")

    # 2. daily table
    doc.add_heading("二、每日营养概览（估算值）", 1)
    tbl = doc.add_table(rows=1, cols=6); tbl.style = "Light Grid Accent 1"
    for i, h in enumerate(["日期", "热量 kcal", "蛋白 g", "碳水 g", "脂肪 g", "备注"]):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for d in dates:
        e = days[d]; note = []
        if "breakfast" not in e["meals"]: note.append("无早餐")
        if len(e["meals"]) < 2 and e["items"] < 2: note.append("仅单餐·未记全")
        if e["f"] >= FAT_NOTE_G: note.append(f"脂肪超{FAT_NOTE_G}g")
        c = tbl.add_row().cells
        c[0].text = d; c[1].text = f"{e['cal']:.0f}"; c[2].text = f"{e['p']:.0f}"
        c[3].text = f"{e['c']:.0f}"; c[4].text = f"{e['f']:.0f}"
        c[5].text = "；".join(note)
    c = tbl.add_row().cells
    c[0].paragraphs[0].add_run(f"均值({len(complete)}天)").bold = True
    for i, k in enumerate(["cal", "p", "c", "f"], 1):
        c[i].paragraphs[0].add_run(f"{avg(k):.0f}").bold = True
    c[5].text = "已排除残缺日"
    cap = doc.add_paragraph(
        f"⚠️ 多数缺早餐的日子为低估值；若有进食，真实摄入可能更接近 "
        f"{HIGH_CAL_DAY} kcal/天。多数条目为 low/med 置信度，份量为目测估算。")
    cap.runs[0].font.size = Pt(9); cap.runs[0].italic = True

    # 2b. charts
    doc.add_heading("营养趋势图（红色虚线为示例阈值）", 2)
    doc.add_paragraph("下图仅含记录完整的日子。热量与脂肪用「上限」红线——超线标红；"
                      "蛋白用「下限」红线——低于线标红。阈值为示例值，可在脚本顶部调整。")
    chartdir = tempfile.mkdtemp(prefix="kibble-charts-")
    for key, cfg, flag, png in make_charts(days, complete, chartdir):
        doc.add_picture(png, width=Inches(6.3))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        crossed = [complete[i][5:] for i in range(len(complete)) if flag[i]]
        verb = "超上限" if cfg["kind"] == "ceiling" else "低于下限"
        capp = doc.add_paragraph()
        r = capp.add_run(f"{cfg['label']}：{verb}的日子 = "
                         + ("、".join(crossed) if crossed else "无")
                         + f"（红线 {cfg['line']}）")
        r.font.size = Pt(8); r.italic = True
        capp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 3 + 4. data-driven signals
    good, issues = build_signals(days, dates, complete, rows)
    doc.add_heading("三、数据信号：做得好的地方", 1)
    if good:
        for t, body in good:
            par = doc.add_paragraph(style="List Bullet")
            par.add_run(t + "：").bold = True; par.add_run(body)
    else:
        doc.add_paragraph("（暂无达到阈值的正向信号）", style="List Bullet")
    doc.add_heading("四、数据信号：需要改进", 1)
    if issues:
        for t, body in issues:
            par = doc.add_paragraph(style="List Number")
            par.add_run(t + "：").bold = True; par.add_run(body)
    else:
        doc.add_paragraph("（各项均在阈值内，无触发的问题信号）", style="List Number")

    note = doc.add_paragraph()
    note.add_run("说明：以上信号由脚本按阈值确定性生成，每条都对应一个具体数字；"
                 "未触发阈值的项目不会生成空泛建议。质性解读由对话补充。").italic = True
    note.runs[0].font.size = Pt(9)

    doc.save(out)
    print("WROTE", out)
    print(f"days={len(dates)} complete={len(complete)} entries={len(rows)} "
          f"avg_cal={avg('cal'):.0f} avg_p={avg('p'):.0f} avg_f={avg('f'):.0f} "
          f"breakfast_days={bf_days} fast={fast} sweet={sweet}")
    return out


if __name__ == "__main__":
    if len(sys.argv) < 2:
        sys.exit("usage: generate_diet_report.py <repo_path>")
    build_report(os.path.expanduser(sys.argv[1]))
